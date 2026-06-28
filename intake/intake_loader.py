import os
from pathlib import Path
import json
import pandas as pd
from docx import Document
import fitz #3

from intake.source_detector import SourceDetector

class IntakeLoader:

    def load(self, file_path: str) -> dict:
        """
        Main entry point for loading any supported file type.

        Args:
            file_path (str): Path to the input file.

        Returns:
            dict: Standardized extraction result.
        """

        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(
                f"File not found: {file_path}"
            )

        source_type = SourceDetector.detect(str(file_path))

        loaders = {
            "pdf": self._load_pdf,
            "docx": self._load_docx,
            "txt": self._load_txt,
            "csv": self._load_csv,
            "json": self._load_json,
        }

        extracted_data = loaders[source_type](file_path)

        return {
            "source_type": source_type,
            "metadata": {
                "file_name": file_path.name,
                "file_extension": file_path.suffix,
                "file_size_bytes": file_path.stat().st_size,
            },
            **extracted_data,
        }
    
    #3
    def _load_pdf(self, file_path: Path) -> dict:
        """
        Extract text and metadata from PDF using PyMuPDF.
        """

        document = fitz.open(str(file_path))

        extracted_pages = []

        for page in document:

            page_text = page.get_text("text", sort = True)

            if page_text:
                extracted_pages.append(page_text.strip())

        document.close()

        full_text = "\n\n".join(extracted_pages)

        return {
            "extraction": {
                "status": "success",
                "content": full_text,
                "page_count": len(extracted_pages)
            }
        }
    
    def _load_docx(self, file_path: Path) -> dict:
        """
        Extract text from DOCX.
        """

        document = Document(str(file_path))

        text = "\n".join(
            paragraph.text
            for paragraph in document.paragraphs
        )

        return {
            "extraction": {
                "status": "success",
                "content": text.strip(),
                "paragraph_count": len(document.paragraphs),
            }
        }

    def _load_txt(self, file_path: Path) -> dict:
        """
        Extract text from TXT.
        """

        with open(
            file_path,
            "r",
            encoding="utf-8"
        ) as file:

            text = file.read()

        return {
            "extraction": {
                "status": "success",
                "content": text,
            }
        }

    def _load_csv(self, file_path: Path) -> dict:
        """
        Extract content and metadata from CSV.
        """

        dataframe = pd.read_csv(file_path)

        return {
            "extraction": {
                "status": "success",
                "content": dataframe.to_string(index=False),
                "row_count": len(dataframe),
                "column_count": len(dataframe.columns),
                "columns": dataframe.columns.tolist(),
            }
        }

    def _load_json(self, file_path: Path) -> dict:
        """
        Extract content from JSON.
        """

        with open(
            file_path,
            "r",
            encoding="utf-8"
        ) as file:

            data = json.load(file)

        return {
            "extraction": {
                "status": "success",
                "content": json.dumps(
                    data,
                    indent=4
                ),
            }
        }